"""Build a Word (.docx) version of the memory_assistant cost summary with charts.

Reads data/cost_report_exp005_variability.json, renders variance visuals with
matplotlib, and assembles agent_summaries/memory_assistant.docx.
"""

import json
import os
import sys
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

REPO = Path(__file__).resolve().parents[1]
DATA = REPO / "data"
OUT_DIR = REPO / "agent_summaries"
IMG = DATA / "summary_charts"
IMG.mkdir(parents=True, exist_ok=True)

BLUE, ORANGE, GREEN, RED = "#4285F4", "#FB8C00", "#34A853", "#EA4335"

rep = json.loads((DATA / "cost_report_exp005_variability.json").read_text())
runs = rep["runs"]
var = rep["variability"]
mem = rep["memory_bank"]
rt = rep["runtime"]
n = len(runs)


def chart_cv():
    """Bar chart: coefficient of variation per metric (which usage is volatile)."""
    metrics = ["input_tokens", "output_tokens", "model_calls", "session_events", "model_usd"]
    labels = ["Input\ntokens", "Output\ntokens", "Model\ncalls", "Session\nevents", "Model\ncost"]
    cvs = [var[m]["cv_pct"] for m in metrics]
    colors = [RED if c >= 40 else (ORANGE if c >= 20 else GREEN) for c in cvs]
    fig, ax = plt.subplots(figsize=(6.5, 3.2))
    bars = ax.bar(labels, cvs, color=colors)
    ax.set_ylabel("Coefficient of variation (%)")
    ax.set_title("Run-to-run variability by metric (higher = noisier)")
    for b, c in zip(bars, cvs):
        ax.text(b.get_x() + b.get_width() / 2, c + 1, f"{c:.0f}%", ha="center", fontsize=9)
    ax.set_ylim(0, max(cvs) * 1.2)
    ax.axhline(40, ls="--", lw=0.8, color="gray")
    fig.tight_layout()
    p = IMG / "cv.png"; fig.savefig(p, dpi=150); plt.close(fig)
    return p


def chart_per_run_cost():
    """Per-run model cost bars + mean line (shows the spread)."""
    vals = [r["model_usd"] for r in runs]
    mean = var["model_usd"]["mean"]
    fig, ax = plt.subplots(figsize=(6.5, 3.2))
    bars = ax.bar([f"Run {i+1}" for i in range(n)], vals, color=BLUE)
    ax.axhline(mean, color=RED, ls="--", label=f"mean ${mean:.4f}")
    for b, v in zip(bars, vals):
        ax.text(b.get_x() + b.get_width() / 2, v, f"${v:.4f}", ha="center", va="bottom", fontsize=8)
    ax.set_ylabel("Model cost per interaction ($)")
    ax.set_title(f"Model cost per run — {max(vals)/min(vals):.1f}x spread, identical task")
    ax.legend()
    fig.tight_layout()
    p = IMG / "per_run_cost.png"; fig.savefig(p, dpi=150); plt.close(fig)
    return p


def chart_cost_breakdown():
    """Typical per-interaction cost by SKU (horizontal bar)."""
    conv = var["model_usd"]["mean"]
    gen = mem["generate_memories_usd"] / n
    retr = mem["memories_retrieved_usd"] / n
    sess = var["session_events"]["mean"] * 0.00025
    runtime = rt["runtime_total_usd"] / n
    items = [("Conversation tokens", conv), ("Agent Runtime", runtime),
             ("Session events", sess), ("Memory retrievals", retr),
             ("Memory generation", gen)]
    items.sort(key=lambda x: x[1])
    names = [i[0] for i in items]; vals = [i[1] for i in items]
    total = sum(vals)
    fig, ax = plt.subplots(figsize=(6.5, 3.2))
    bars = ax.barh(names, vals, color=[GREEN, BLUE, ORANGE, "#9C27B0", "#00ACC1"])
    for b, v in zip(bars, vals):
        ax.text(v, b.get_y() + b.get_height() / 2, f" ${v:.4f} ({v/total:.0%})",
                va="center", fontsize=8)
    ax.set_xlabel("USD per interaction (catalog list price)")
    ax.set_title(f"Cost by SKU — total ≈ ${total:.4f}/interaction")
    ax.set_xlim(0, max(vals) * 1.35)
    fig.tight_layout()
    p = IMG / "breakdown.png"; fig.savefig(p, dpi=150); plt.close(fig)
    return p, total


def H(doc, text, level):
    doc.add_heading(text, level=level)


def main():
    cv_png = chart_cv()
    cost_png = chart_per_run_cost()
    bd_png, total = chart_cost_breakdown()

    doc = Document()
    title = doc.add_heading("Agent Cost Summary — memory_assistant", level=0)

    p = doc.add_paragraph()
    p.add_run("Model: ").bold = True; p.add_run("gemini-2.5-flash    ")
    p.add_run("Platform: ").bold = True; p.add_run("Vertex AI Agent Engine (GEAP Agent Runtime)")
    p2 = doc.add_paragraph()
    p2.add_run("Experiments: ").bold = True
    p2.add_run("EXP-004 (full priced breakdown), EXP-005 (4-run variability)")
    p3 = doc.add_paragraph()
    p3.add_run("Cost unit: ").bold = True
    p3.add_run("1 interaction = 3 user messages across 2 sessions (2 facts + 1 recall) + "
               "memory generation; ~5.75 model calls. Not comparable to single-query agents.")

    H(doc, "1. Architecture", 1)
    doc.add_paragraph(
        "A coordinator with long-term memory delegating to two specialist sub-agents:")
    for line in [
        "personal_assistant (coordinator) — tool: preload_memory",
        "  └─ prefs_agent  → set_unit_preference, convert_temp",
        "  └─ notes_agent  → make_checklist",
        "Memory Bank (auto-wired on deploy) — generate / retrieve / store",
        "Sessions (managed, persistent) — every turn/event persisted",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    H(doc, "2. Components / SKUs used", 1)
    comp = [("Gemini 2.5 Flash", "input/output/cached tokens", "usage_metadata"),
            ("Agent Runtime", "vCPU-sec + GiB-sec", "Cloud Monitoring allocation_time"),
            ("Memory Bank — generate", "Gemini tokens (server-side)", "Monitoring generate_memories_token_count"),
            ("Memory Bank — retrieve", "per retrieval op", "Monitoring memory_retrieval_count"),
            ("Memory Bank — store", "per memory / month", "export-only"),
            ("Sessions", "per event appended", "observed events (approx)")]
    t = doc.add_table(rows=1, cols=3); t.style = "Light Grid Accent 1"
    for i, h in enumerate(["Component", "Meter", "Capture source"]):
        t.rows[0].cells[i].paragraphs[0].add_run(h).bold = True
    for c in comp:
        cells = t.add_row().cells
        for i, v in enumerate(c):
            cells[i].text = v

    H(doc, "3. How the experiments were run", 1)
    for s in [
        "Deploy once to Agent Engine; Memory Bank + managed Sessions auto-wire.",
        "Per run: Session A (2 fact turns) → add_session_to_memory → 20s wait → Session B (recall).",
        "Capture per-run token usage instantly from usage_metadata.",
        "Settle ~300s for Monitoring ingestion, then pull runtime + memory_bank metrics (60s alignment).",
        "Repeat N times (EXP-005: 4 runs, fresh user each) → mean / CV% / min–max per dimension.",
        "Price every captured quantity × live Billing Catalog rate.",
    ]:
        doc.add_paragraph(s, style="List Number")
    doc.add_paragraph("Reproduce: python scripts/exp005_variability.py --runs 4 --settle 300")

    H(doc, "4. Typical usage & variance", 1)
    rows_tbl = [("Metric", "mean", "min–max", "CV%")]
    for key, lab in [("input_tokens", "Input tokens"), ("output_tokens", "Output tokens (incl. thinking)"),
                     ("model_calls", "Model calls"), ("session_events", "Session events"),
                     ("model_usd", "Model cost ($)")]:
        v = var[key]
        rows_tbl.append((lab, f"{v['mean']}", f"{v['min']}–{v['max']}", f"{v['cv_pct']}%"))
    t2 = doc.add_table(rows=0, cols=4); t2.style = "Light Grid Accent 1"
    for r in rows_tbl:
        cells = t2.add_row().cells
        for i, v in enumerate(r):
            run = cells[i].paragraphs[0].add_run(v)
            if r is rows_tbl[0]:
                run.bold = True
    doc.add_paragraph(f"Recall success rate: {rep['recall_rate']:.0%} (function reliable; cost is what varies).")

    doc.add_picture(str(cv_png), width=Inches(6.0))
    doc.add_picture(str(cost_png), width=Inches(6.0))
    doc.add_picture(str(bd_png), width=Inches(6.0))

    H(doc, "5. Summary", 1)
    sm = doc.add_paragraph()
    sm.add_run(f"Budget ≈ ${total:.4f} per interaction at list price").bold = True
    sm.add_run(", but treat it as a distribution: model-token cost swings ~"
               f"{max(r['model_usd'] for r in runs)/min(r['model_usd'] for r in runs):.1f}x "
               "run-to-run (CV 48%), driven by output/thinking tokens (CV 57%). Structural usage "
               "(model calls, session events, input tokens) is stable (CV 8–16%). Runtime cost is "
               "utilization-dependent — pin a queries/hour assumption. For an SLA/quote use the "
               "p50–max range, not the mean.")

    H(doc, "6. Caveats", 1)
    for c in ["Catalog list price, not the project's actual billed rate (needs BigQuery export).",
              "Memory storage (monthly) and authoritative session-event counts are export-only.",
              "Cloud Trace (enabled on deploy), Logging, GCS, egress not captured.",
              "Runtime amortization depends on the measurement window's utilization."]:
        doc.add_paragraph(c, style="List Bullet")

    out = OUT_DIR / "memory_assistant.docx"
    doc.save(str(out))
    print("Wrote", out)
    print(f"Charts: {cv_png.name}, {cost_png.name}, {bd_png.name}")
    print(f"Total per interaction: ${total:.4f}")


if __name__ == "__main__":
    main()
